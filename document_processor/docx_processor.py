# [file name]: docx_processor.py
import os
import re
from pathlib import Path
from typing import List, Dict, Any

from base_processor import BaseDocumentProcessor

# Импорты для DOCX
try:
    from docx import Document
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False


class DOCXProcessor(BaseDocumentProcessor):
    """Процессор для DOCX файлов с улучшенной разметкой"""

    def can_process(self, file_path: str) -> bool:
        return file_path.lower().endswith('.docx') and HAS_DOCX

    def extract_content_with_tables_ordered(self, file_path: str) -> str:
        """Извлечение содержимого DOCX с улучшенной разметкой"""
        try:
            doc = Document(file_path)
            content_lines = []

            for element in self._get_document_elements_ordered(doc):
                if element['type'] == 'paragraph':
                    paragraph = element['object']
                    if paragraph.text.strip():
                        formatted_text = self._format_paragraph_improved(paragraph)
                        if formatted_text:
                            content_lines.append(formatted_text)

                elif element['type'] == 'table':
                    table = element['object']
                    if self._has_table_content(table):
                        table_text = self._format_table_universal(table)
                        if table_text:
                            content_lines.append(table_text)

            return '\n'.join(content_lines)

        except Exception as e:
            self.logger.error(f"Ошибка извлечения содержимого из DOCX {file_path}: {e}")
            raise

    def _format_paragraph_improved(self, paragraph) -> str:
        """Улучшенное форматирование параграфа"""
        text = paragraph.text.strip()
        if not text:
            return ""

        # 1. Определяем уровень заголовка по стилю
        if paragraph.style and paragraph.style.name:
            style_name = paragraph.style.name.lower()
            
            if any(heading in style_name for heading in ['heading 1', 'title', 'heading1', 'заголовок 1']):
                return f"# {text}"
            elif any(heading in style_name for heading in ['heading 2', 'heading2', 'заголовок 2']):
                return f"## {text}"
            elif any(heading in style_name for heading in ['heading 3', 'heading3', 'заголовок 3']):
                return f"### {text}"
            elif any(heading in style_name for heading in ['heading 4', 'heading4', 'заголовок 4']):
                return f"#### {text}"

        # 2. Проверяем, если весь абзац выделен жирным
        if self._is_fully_bold(paragraph):
            # Определяем уровень заголовка по длине текста и контексту
            if len(text) < 100:  # Короткий жирный текст - вероятно заголовок
                if len(text) < 50:
                    return f"## {text}"  # Короткий заголовок - уровень 2
                else:
                    return f"### {text}"  # Более длинный - уровень 3

        # Обычный текст без форматирования заголовков
        return text

    def _is_fully_bold(self, paragraph) -> bool:
        """Проверяет, является ли весь абзац жирным"""
        if not paragraph.runs:
            return False
        
        # Проверяем, что все runs с текстом являются жирными
        for run in paragraph.runs:
            if run.text.strip() and not run.bold:
                return False
        
        return any(run.text.strip() for run in paragraph.runs)

    def _format_table_universal(self, table) -> str:
        """Универсальное форматирование таблицы с шапкой 0-3 строк"""
        table_lines = []

        # Маркер начала таблицы ДО шапки
        table_lines.append("--- Таблица начало ---")
        
        # Извлекаем содержимое таблицы
        table_data = []
        max_columns = 0

        for row in table.rows:
            row_cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                cell_text = re.sub(r'\s+', ' ', cell_text)
                row_cells.append(cell_text)

            if row_cells and any(cell_text for cell_text in row_cells):
                table_data.append(row_cells)
                max_columns = max(max_columns, len(row_cells))

        if not table_data:
            return ""

        # 1. Автоматическое определение количества строк шапки
        header_row_count = self._detect_header_rows(table_data)
        
        # 2. Разделение на шапку и данные
        header_rows = table_data[:header_row_count] if header_row_count > 0 else []
        data_rows = table_data[header_row_count:]
        
        # 3. Объединение шапки (если есть)
        if header_rows:
            merged_header = self._merge_header_columns(header_rows)
            if merged_header:
                table_lines.append(" | ".join(merged_header))
                # Добавляем разделитер
                separator = ['---'] * len(merged_header)
                table_lines.append(" | ".join(separator))
        
        # 4. Добавление данных
        for row in data_rows:
            # Дополняем строки до максимального количества колонок
            while len(row) < max_columns:
                row.append("")
            table_lines.append(" | ".join(row))

        # Маркер конца таблицы
        table_lines.append("--- Таблица конец ---")
        table_lines.append("")  # Пустая строка после таблицы

        return "\n".join(table_lines)

    def _detect_header_rows(self, table_data: List[List[str]]) -> int:
        """Автоматически определяет количество строк шапки (0-3)"""
        if not table_data:
            return 0
        
        max_header_rows = min(3, len(table_data))
        
        for i in range(max_header_rows):
            row = table_data[i]
            
            # Критерии для определения конца шапки:
            # 1. Следующая строка содержит числовые данные
            if i + 1 < len(table_data) and self._has_numeric_data(table_data[i + 1]):
                return i + 1
            
            # 2. Текущая строка выглядит как данные (длинный текст, числа)
            if self._looks_like_data_row(row):
                return i
            
            # 3. Текущая строка - явная шапка (короткий текст, заголовочные слова)
            if self._is_explicit_header_row(row):
                continue
                
            # 4. По умолчанию считаем первую строку шапкой
            if i == 0:
                continue
                
        return max_header_rows

    def _has_numeric_data(self, row: List[str]) -> bool:
        """Проверяет, содержит ли строка числовые данные"""
        for cell in row:
            if cell and re.search(r'\d+[,\.]\d+|\b\d+\b', str(cell)):
                return True
        return False

    def _looks_like_data_row(self, row: List[str]) -> bool:
        """Проверяет, выглядит ли строка как строка данных"""
        if not row:
            return False
        
        # Данные обычно содержат более длинный текст или числа
        total_chars = sum(len(str(cell)) for cell in row if cell)
        avg_chars_per_cell = total_chars / len(row)
        
        return avg_chars_per_cell > 25 or self._has_numeric_data(row)

    def _is_explicit_header_row(self, row: List[str]) -> bool:
        """Проверяет, является ли строка явной шапкой"""
        if not row:
            return False
        
        header_keywords = ['характеристика', 'показатель', 'параметр', 'наименование', 
                          'описание', 'тип', 'вид', 'категория', 'название']
        
        row_text = ' '.join(str(cell).lower() for cell in row if cell)
        has_header_keywords = any(keyword in row_text for keyword in header_keywords)
        
        # Шапка обычно короткая
        total_chars = sum(len(str(cell)) for cell in row if cell)
        avg_chars_per_cell = total_chars / len(row) if row else 0
        
        return has_header_keywords and avg_chars_per_cell < 30

    def _merge_header_columns(self, header_rows: List[List[str]]) -> List[str]:
        """Объединяет шапку по колонкам (вертикальное объединение)"""
        if not header_rows:
            return []
            
        if len(header_rows) == 1:
            return header_rows[0]
            
        # Определяем максимальное количество колонок
        max_cols = max(len(row) for row in header_rows)
        
        # Создаем объединенную шапку по колонкам
        merged_header = []
        
        for col_idx in range(max_cols):
            column_cells = []
            for row in header_rows:
                if col_idx < len(row) and row[col_idx]:
                    cell_text = str(row[col_idx]).strip()
                    if cell_text:
                        column_cells.append(cell_text)
            
            if column_cells:
                # Объединяем ячейки вертикально через пробел
                merged_cell = " ".join(column_cells)
                merged_cell = re.sub(r'\s+', ' ', merged_cell).strip()
                merged_header.append(merged_cell)
            else:
                merged_header.append("")
                
        return merged_header

    def _get_document_elements_ordered(self, doc) -> List[Dict]:
        """Получает все элементы документа в правильном порядке"""
        elements = []

        try:
            for block in doc.element.body:
                if block.tag.endswith('p'):  # Параграф
                    try:
                        paragraph = self._get_paragraph_from_element(block, doc)
                        if paragraph:
                            elements.append({'type': 'paragraph', 'object': paragraph})
                    except Exception as e:
                        self.logger.warning(f"Ошибка создания параграфа: {e}")

                elif block.tag.endswith('tbl'):  # Таблица
                    try:
                        table = self._get_table_from_element(block, doc)
                        if table:
                            elements.append({'type': 'table', 'object': table})
                    except Exception as e:
                        self.logger.warning(f"Ошибка создания таблицы: {e}")
        except Exception as e:
            self.logger.warning(f"Ошибка при разборе документа: {e}")
            return self._get_document_elements_fallback(doc)

        return elements

    def _get_document_elements_fallback(self, doc) -> List[Dict]:
        """Резервный метод получения элементов документа"""
        elements = []

        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                elements.append({'type': 'paragraph', 'object': paragraph})

        for table in doc.tables:
            if self._has_table_content(table):
                elements.append({'type': 'table', 'object': table})

        return elements

    def _get_paragraph_from_element(self, element, doc):
        """Создает объект Paragraph из XML элемента"""
        from docx.text.paragraph import Paragraph
        return Paragraph(element, doc)

    def _get_table_from_element(self, element, doc):
        """Создает объект Table из XML элемента"""
        from docx.table import Table
        return Table(element, doc)

    def _has_table_content(self, table) -> bool:
        """Проверяет, содержит ли таблица полезное содержимое"""
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    return True
        return False

    def convert_to_txt(self, file_path: str, output_dir: str) -> str:
        """Конвертация DOCX в TXT"""
        try:
            self.logger.info(f"Конвертация DOCX в TXT: {file_path}")

            text = self.extract_content_with_tables_ordered(file_path)

            output_path = os.path.join(output_dir, f"{Path(file_path).stem}.txt")
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(text)

            self.logger.info(f"DOCX конвертирован в: {output_path}")
            return output_path

        except Exception as e:
            self.logger.error(f"Ошибка конвертации DOCX {file_path}: {e}")
            raise